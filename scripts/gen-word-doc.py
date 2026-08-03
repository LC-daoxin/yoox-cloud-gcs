#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《YOOX Cloud GCS 上云 API 开发文档》Word 文件。

内容与在线文档门户（api-portal）一致，基于本项目自身 API 实现与公开的
KMZ/WPML 结构原创编写，供客户开发使用。运行：
    /tmp/yoox-docx-venv/bin/python scripts/gen-word-doc.py
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "Microsoft YaHei"
MONO = "Consolas"
OUT = "docs/YOOX-Cloud-GCS-上云API开发文档.docx"

BRAND = RGBColor(0x1F, 0x6F, 0xB2)
INK = RGBColor(0x22, 0x28, 0x30)
MUTED = RGBColor(0x66, 0x70, 0x7D)
CODE_BG = "F2F4F7"
HEAD_BG = "EAF1F8"
INFO_BG = "EAF3FF"
WARN_BG = "FFF6E6"
DANGER_BG = "FDECEC"


def set_cjk(run, font=CJK):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:eastAsia'), font)


def shade(el, hex_color):
    pr = el.get_or_add_tcPr() if el.tag.endswith('}tc') else el.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pr.append(shd)


def cell_bg(cell, hex_color):
    shade(cell._tc, hex_color)


def para_shade(p, hex_color):
    shade(p._p, hex_color)


def add_lead(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    set_cjk(r)
    p.space_after = Pt(6)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    set_cjk(r)
    return p


def add_bullets(doc, items, ordered=False):
    style = 'List Number' if ordered else 'List Bullet'
    for it in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(it)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
        set_cjk(r)


def add_code(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell_bg(cell, CODE_BG)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x33, 0x4D)
        rpr = r._element.get_or_add_rPr()
        rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), MONO); rf.set(qn('w:hAnsi'), MONO); rpr.append(rf)
    _thin_borders(tbl, "D0D7DE")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, kind, label, text):
    bg = {"info": INFO_BG, "warn": WARN_BG, "danger": DANGER_BG}[kind]
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell_bg(cell, bg)
    p = cell.paragraphs[0]
    rl = p.add_run(label + "：")
    rl.bold = True; rl.font.size = Pt(10); rl.font.color.rgb = INK; set_cjk(rl)
    rt = p.add_run(text)
    rt.font.size = Pt(10); rt.font.color.rgb = INK; set_cjk(rt)
    _thin_borders(tbl, "E2C98A" if kind != "info" else "AECBF0")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell_bg(hdr[i], HEAD_BG)
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK; set_cjk(r)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            mono = val.startswith("`") and val.endswith("`")
            text = val.strip("`")
            r = p.add_run(text)
            if mono:
                r.font.name = MONO; r.font.size = Pt(8.5)
                rpr = r._element.get_or_add_rPr()
                rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), MONO); rf.set(qn('w:hAnsi'), MONO); rpr.append(rf)
            else:
                r.font.size = Pt(9.5); set_cjk(r)
            r.font.color.rgb = INK
    _thin_borders(tbl, "D0D7DE")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _thin_borders(tbl, color):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def h1(doc, text):
    doc.add_page_break()
    p = doc.add_heading(level=1)
    r = p.add_run(text); set_cjk(r); r.font.color.rgb = BRAND
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); set_cjk(r); r.font.color.rgb = INK
    return p


def h3(doc, text):
    p = doc.add_heading(level=3)
    r = p.add_run(text); set_cjk(r); r.font.color.rgb = INK
    return p


def build():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = CJK
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), CJK); rf.set(qn('w:hAnsi'), CJK); rf.set(qn('w:eastAsia'), CJK)
    rpr.append(rf)

    # ---------- 封面 ----------
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("YOOX Cloud GCS"); rt.bold = True; rt.font.size = Pt(30); rt.font.color.rgb = BRAND; set_cjk(rt)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("上云 API 开发文档"); rs.font.size = Pt(20); rs.font.color.rgb = INK; set_cjk(rs)
    doc.add_paragraph()
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("面向客户开发的设备上云协议与云端接口说明"); rsub.font.size = Pt(12); rsub.font.color.rgb = MUTED; set_cjk(rsub)
    for _ in range(8):
        doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"版本 1.0 · 生成日期 {datetime.date.today().isoformat()}"); rm.font.size = Pt(11); rm.font.color.rgb = MUTED; set_cjk(rm)
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("内部分发 · 本文内容基于本项目 API 实现与公开的 KMZ/WPML 结构原创编写")
    rn.font.size = Pt(9); rn.font.color.rgb = MUTED; set_cjk(rn)

    # ---------- 目录 ----------
    doc.add_page_break()
    ph = doc.add_heading(level=1); rph = ph.add_run("目录"); set_cjk(rph); rph.font.color.rgb = BRAND
    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    ttxt = OxmlElement('w:t'); ttxt.text = "在 Word 中按 F9 或右键“更新域”生成目录…"
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep); run._r.append(ttxt); run._r.append(fldEnd)

    render(doc)
    import os
    os.makedirs("docs", exist_ok=True)
    doc.save(OUT)
    print("已生成:", OUT)


def render(doc):
    # ===== 1. 基础介绍 =====
    h1(doc, "1. 基础介绍")
    add_lead(doc, "YOOX Cloud GCS 是一套可容器化私有部署的网页云端地面站，基于设备上云协议为行业无人机、遥控器与机巢提供设备管理、实时直播、远程控制、航线任务与媒体归档能力。")
    h2(doc, "1.1 平台简介")
    add_para(doc, "YOOX Cloud GCS 面向行业无人机的云端统一接入与作业场景，核心能力包括：")
    add_bullets(doc, [
        "设备管理：网关（遥控器/机巢）与子设备（飞行器/负载）的绑定、拓扑、属性与在线状态。",
        "实时直播：设备侧以 RTSP 推流，服务端经 MediaMTX 转封装为 WebRTC，浏览器通过 WHEP 会话播放。",
        "虚拟座舱 / 远程控制：DRC（Direct Remote Control）指令飞行，含控制权租约、心跳与安全联锁。",
        "航线任务：上传 KMZ 航线文件、创建与调度飞行任务、跟踪进度并回传媒体。",
        "媒体归档：媒体文件入对象存储，短时授权 URL 访问。",
    ])
    add_callout(doc, "info", "协议边界", "上云协议的推流能力矩阵包含 RTMP、RTSP、GB28181。当前运行版本仅开放 RTSP 入流，其余为后期能力。")
    h2(doc, "1.2 系统架构")
    add_para(doc, "平台采用前后端分离 + 容器编排。设备通过 MQTT 接入消息代理，Web 控制台与后端服务通过 HTTP / WebSocket 交互，媒体与对象存储独立成栈。")
    add_code(doc, "设备/机巢 ──MQTT──▶ EMQX ◀──▶ Cloud Service(Spring Boot) ◀──▶ Web 控制台(Vue3)\n设备 RTSP ──▶ MediaMTX ──WHEP──▶ 浏览器 WebRTC   |   MinIO 对象存储")
    add_table(doc, ["组件", "职责"], [
        ["Cloud Service", "业务 API、设备物模型、鉴权、DRC 控制、航线任务编排"],
        ["EMQX", "设备与云端 MQTT 消息代理，基于 Redis 的 Topic ACL"],
        ["MediaMTX", "RTSP 入流转 WebRTC（WHEP）出流"],
        ["MySQL / Redis", "业务持久化 / 会话、缓存、OSD 与 ACL"],
        ["MinIO", "航线与媒体对象存储，签发短时访问凭据"],
    ])
    h2(doc, "1.3 接口通用约定")
    add_para(doc, "HTTP API 统一由 Web 入口反向代理，默认基地址 BASE_URL=http://<host>:<http-port>。除登录、令牌刷新、健康检查与 OpenAPI 外，请求须携带鉴权头：")
    add_code(doc, "x-auth-token: <access_token>\nContent-Type: application/json")
    add_para(doc, "统一响应信封：")
    add_code(doc, '{\n  "code": 0,\n  "message": "success",\n  "data": {}\n}')
    add_para(doc, "code=0 表示业务成功；HTTP 401 或业务 code=401 表示令牌无效。调用方必须同时检查 HTTP 状态码与业务 code。接口按业务域划分前缀：manage、map、media、wayline、storage、control，版本段统一为 /api/v1。")
    h2(doc, "1.4 登录与鉴权")
    add_para(doc, "账号分两类：flag=1 为 Web 账号，flag=2 为 Pilot（遥控器）账号。登录成功返回 access_token、workspace_id 及用户信息。")
    add_code(doc, "curl -sS \"$BASE_URL/manage/api/v1/login\" \\\n  -H 'content-type: application/json' \\\n  -d '{\"username\":\"admin\",\"password\":\"<password>\",\"flag\":1}'")
    add_callout(doc, "info", "令牌机制", "令牌为 HMAC-SHA256 签名的 JWT，默认有效期 86400 秒。签名密钥由部署侧通过环境变量注入；修改密钥会使既有 Web 与 DRC 会话立即失效。")
    h2(doc, "1.5 快速开始")
    add_bullets(doc, [
        "调用登录接口获取 access_token 与 workspace_id。",
        "拉取设备拓扑：GET /manage/api/v1/devices/{workspace_id}/devices。",
        "查询直播能力并启动 RTSP 视频源，前端用返回的 WHEP URL 播放。",
        "如需远控，进入 DRC 流程获取控制权与短时 MQTT 凭据。",
        "上传 KMZ 航线并创建飞行任务，通过 WebSocket 跟踪进度。",
    ], ordered=True)

    # ===== 2. 产品与协议介绍 =====
    h1(doc, "2. 产品与协议介绍")
    add_lead(doc, "本章介绍设备上云协议的定位与核心理念、端边云分层架构、设备域与产品标识规则，以及平台开放的功能集合总览。")
    h2(doc, "2.1 上云协议概述")
    add_para(doc, "传统做法中，无人机接入第三方云平台通常需要基于移动端 SDK 定制 App 并自定义私有通信协议，开发者需要在飞行功能接口适配上投入大量精力。设备上云协议将无人机、遥控器、机巢的能力抽象为物联网设备的物模型，开发者无需重新开发 App，即可把设备接入云平台，从而专注于上云业务本身。")
    h2(doc, "2.2 核心理念")
    add_bullets(doc, [
        "标准协议：采用业界通用的 MQTT、HTTPS、WebSocket 协议承载消息、请求与实时推送。",
        "物模型抽象：把设备能力抽象为属性（state）、事件（events）、服务（services）三类，开发者面向物模型开发。",
        "能力驱动：所有功能均以设备上报的能力为准动态开放，前端按能力渲染，避免下发未验证指令。",
        "部署无关：只要通信链路能访问到第三方云平台，私有化或公有云部署均可工作。",
    ])
    h2(doc, "2.3 端边云分层架构")
    add_para(doc, "上云协议在设备端 SDK 之上封装为独立功能集，采用与物联网类似的「端—边—云」分层。网关设备（遥控器/机巢）与云平台之间使用 MQTT/HTTPS/WebSocket 通信，在协议之上抽象出各硬件设备的物模型及业务功能集。")
    add_code(doc, "端：飞行器/负载 ──▶ 边：网关(遥控器/机巢) ──MQTT/HTTPS/WS──▶ 云：第三方云平台")
    h2(doc, "2.4 设备域与产品标识")
    add_para(doc, "平台通过 domain、type、sub_type 唯一确定一款设备；对于负载，再通过 gimbalindex 确定其挂载于哪款无人机的哪个云台口。")
    add_table(doc, ["字段", "含义"], [
        ["`domain`", "设备领域命名空间：0 无人机类、1 负载类、2 遥控器类、3 机巢类、60 基站类、70 无人车类"],
        ["`type`", "设备主类型"],
        ["`sub_type`", "设备子类型"],
        ["`gimbalindex`", "负载云台口：0 主云台、1 机身下方右云台、2 机身上方云台、7 FPV 相机，其余为预留"],
    ])
    add_callout(doc, "info", "说明", "具体可接入的机型、负载与机巢型号以部署时的产品支持清单与设备实际上报的能力为准。平台按 domain/type/sub_type 组织物模型与拓扑。")
    h2(doc, "2.5 功能集合总览")
    add_table(doc, ["功能集", "能力（通道）"], [
        ["设备与云端连接", "设备绑定信息获取、绑定码绑定组织、绑定组织信息查询（MQTT/HTTP）"],
        ["设备管理", "设备属性推送、拓扑更新、属性设置（MQTT）"],
        ["态势感知", "设备拓扑列表、OSD 遥测定频推送、上下线推送（HTTP/WebSocket）"],
        ["直播功能", "直播能力更新、开始/停止直播、设置清晰度、切换镜头（MQTT/HTTP）"],
        ["地图元素", "点线面创建/更新/删除、workspace 内同步（HTTP/WebSocket）"],
        ["媒体管理", "上传临时凭证、上传结果上报（MQTT）"],
        ["航线管理", "任务下发/执行/取消、进度与状态上报、任务资源获取（MQTT/HTTP）"],
        ["HMS 健康管理", "健康告警上报与文案渲染（MQTT）"],
        ["指令飞行", "DRC 飞行控制、负载控制、flyto、一键起飞（MQTT/HTTP）"],
        ["固件升级（规划）", "固件升级下发与进度（MQTT）"],
        ["远程日志（规划）", "可上传文件列表、发起上传、进度与状态（MQTT/HTTP）"],
    ])

    # ===== 3. 设备上云接入 =====
    h1(doc, "3. 设备上云接入")
    add_lead(doc, "本章说明设备如何通过 MQTT 物模型接入平台，包含接入架构、应用凭据、Topic 体系、消息信封与设备上线数据流。")
    h2(doc, "3.1 接入架构")
    add_para(doc, "设备分为「网关设备」与「子设备」：网关（遥控器或机巢）直接建立 MQTT 连接，子设备（飞行器、负载）经由网关代理上云。平台按网关序列号（gateway_sn）组织 Topic。")
    h2(doc, "3.2 上云交互时序")
    add_code(doc, "1. 设备端 POST /manage/api/v1/login → 返回 access_token、workspace_id\n2. 设备端发起 MQTT 连接（账号/密码），成功后订阅相关 Topic\n3. 设备端发送 update_topo 上报设备拓扑\n4. 平台建立设备拓扑，经 WebSocket 向 Web 端推送 device_update_topo\n5. 设备持续上报 state（属性/能力）与 OSD（遥测）\n6. Web 端建立 WebSocket 连接，接收上下线、OSD、告警、任务进度等推送")
    h2(doc, "3.3 应用凭据")
    add_table(doc, ["配置项", "说明"], [
        ["`APP_ID`", "上云开发者应用 ID"],
        ["`APP_KEY`", "上云开发者应用 Key"],
        ["`APP_LICENSE`", "上云应用 License"],
    ])
    add_callout(doc, "info", "凭据一致性", "这三个值需与设备端使用的开发者应用一致。修改后需重启后端服务使配置生效。")
    h2(doc, "3.4 MQTT Topic 体系")
    add_code(doc, "sys/product/{gateway_sn}/status            # 上下线状态\nthing/product/{gateway_sn}/requests        # 设备发起的请求\nthing/product/{gateway_sn}/requests_reply  # 平台对请求的应答\nthing/product/{gateway_sn}/state           # 设备状态 / 属性\nthing/product/{gateway_sn}/services        # 平台下发的服务调用\nthing/product/{gateway_sn}/services_reply  # 设备对服务调用的应答\nthing/product/{gateway_sn}/events          # 事件上报\nthing/product/{gateway_sn}/events_reply    # 平台对事件的应答\nthing/product/{gateway_sn}/property/set    # 属性设置下发")
    h2(doc, "3.5 消息信封")
    add_para(doc, "物模型消息统一采用如下信封，请求与应答必须关联同一事务 ID（tid）：")
    add_code(doc, '{\n  "tid": "<事务 id>",\n  "bid": "<业务 id>",\n  "timestamp": 1785319200000,\n  "method": "<方法名>",\n  "data": { }\n}')
    add_table(doc, ["字段", "含义"], [
        ["`tid`", "事务 ID，用于请求-应答关联"],
        ["`bid`", "业务 ID，标识一次业务流程"],
        ["`timestamp`", "毫秒级时间戳"],
        ["`method`", "方法名，标识消息语义"],
        ["`data`", "业务数据载荷"],
    ])
    h2(doc, "3.6 设备上线数据流")
    add_bullets(doc, [
        "上线：网关在 sys/product/{gateway_sn}/status 上报上线，平台建立设备拓扑。",
        "属性/状态：设备持续在 state 上报设备属性与能力（如相机/码流能力）。",
        "OSD：飞行器高频上报姿态、位置、电量等遥测（缓存于 Redis）。",
        "事件：HMS 告警、任务进度等经 events 上报。",
        "服务调用：平台下发服务指令，设备经 services_reply 应答。",
    ], ordered=True)

    # ===== 4. 态势感知 =====
    h1(doc, "4. 态势感知")
    add_lead(doc, "态势感知把同一工作空间下所有设备（无人机/遥控器/机巢）的位置与状态形成一张网，服务端定频向各终端推送设备遥测。")
    h2(doc, "4.1 功能概述")
    add_para(doc, "设备上报遥测信息给服务端，服务端定频推送同一工作空间下所有设备的遥测给各终端；终端据此实时更新地图中的设备状态与位置。设备上线/下线或拓扑变化时，服务端广播拓扑更新推送，终端收到后重新拉取设备拓扑列表。")
    h2(doc, "4.2 设备拓扑")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices`", "设备拓扑列表"],
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices/{device_sn}`", "设备详情"],
        ["GET", "`/map/api/v1/workspaces/{workspace_id}/device-status`", "设备状态（地图态势）"],
    ])
    h2(doc, "4.3 OSD 遥测")
    add_table(doc, ["biz_code", "说明"], [
        ["`device_osd`", "飞行器遥测：经纬度、高度、姿态、速度、电量等"],
        ["`dock_osd`", "机巢遥测：舱盖、环境、网络、存储等"],
        ["`gateway_osd`", "网关（遥控器/机巢）遥测"],
        ["`device_update_topo`", "设备上线/下线/拓扑更新推送"],
    ])
    add_callout(doc, "info", "坐标系", "态势与地图元素的经纬度坐标均采用 WGS84 坐标系。前端如使用高德底图，需按需做坐标系转换。")

    # ===== 5. 直播功能 =====
    h1(doc, "5. 直播功能")
    add_lead(doc, "直播功能把无人机相机负载的视频码流发给云平台播放，支持直播的开始、停止、清晰度设置与镜头切换。")
    h2(doc, "5.1 功能概述")
    add_para(doc, "无人机飞行平台不直接连接云平台，中间由网关设备转流转发；遥控器与无人机之间仍走私有图传链路。云平台需预先部署 MQTT 网关与流媒体服务器。")
    add_table(doc, ["直播类型", "说明"], [
        ["`RTSP`", "实时流传输协议，当前运行版本开放的入流方式"],
        ["`RTMP`（规划）", "基于 TCP 的实时消息传输协议族"],
        ["`GB28181`（规划）", "安防视频联网传输控制规范，可对接已有 28181 平台"],
    ])
    h2(doc, "5.2 媒体链路")
    add_code(doc, "设备 RTSP 推流 ──▶ MediaMTX ──WHEP──▶ 浏览器 WebRTC")
    add_para(doc, "设备直播能力（live_capacity）随物模型状态在网关设备上报，仅在设备端状态变化时推送；其中包含可用于直播的视频流总数、可同时直播的视频流总数与设备直播能力列表。")
    h2(doc, "5.3 交互流程")
    add_bullets(doc, [
        "Web 端点击直播；云端解析设备上报的 live_capacity。",
        "下发「开始直播」，指定协议类型与直播质量，设备向流媒体服务器推流。",
        "浏览器用返回的 WHEP URL 拉流播放。",
        "可在不影响直播进程的情况下设置清晰度或切换镜头。",
        "下发「停止直播」结束推流。",
    ], ordered=True)
    h2(doc, "5.4 直播 API")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/manage/api/v1/live/capacity`", "查询直播能力"],
        ["POST", "`/manage/api/v1/live/streams/start`", "开始直播"],
        ["POST", "`/manage/api/v1/live/streams/stop`", "停止直播"],
        ["POST", "`/manage/api/v1/live/streams/update`", "设置直播清晰度"],
        ["POST", "`/manage/api/v1/live/streams/switch`", "切换直播镜头"],
    ])
    add_para(doc, "对应设备侧物模型服务方法：live_start_push、live_stop_push、live_set_quality、live_lens_change；能力状态方法为 live_capacity。清晰度与镜头类型枚举以设备能力上报为准。")

    # ===== 6. 地图元素 =====
    h1(doc, "6. 地图元素")
    add_lead(doc, "地图元素把地图界面上的点、线、面标记同步到云端，实现指挥中心 Web 与设备端飞手之间的实时战术标绘共享。")
    h2(doc, "6.1 功能概述")
    add_para(doc, "用户在 Web 端或设备端绘制点线面时，向服务端发送新增/更新/删除请求；服务端更新存储后，把变更经 WebSocket 通知同一工作空间下的各终端，实现不同飞手之间的元素共享。")
    add_callout(doc, "warn", "注意", "只有正确设置工作空间（workspace）后，才会开始同步该 workspace 下的元素；地图元素坐标采用 WGS84 坐标系；设备端能力可能仅支持打点，线/面视端能力而定。")
    h2(doc, "6.2 元素模型")
    add_para(doc, "元素按图层分组（element_group）管理，元素资源（resource）承载几何与样式信息，并记录创建者。几何类型通常包含点、线、面，坐标以经纬度序列表达。")
    h2(doc, "6.3 HTTP 接口")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/map/api/v1/workspaces/{workspace_id}/element-groups`", "获取图层与元素列表"],
        ["POST", "`/map/api/v1/workspaces/{workspace_id}/element-groups/{group_id}/elements`", "创建地图元素"],
        ["PUT", "`/map/api/v1/workspaces/{workspace_id}/elements/{element_id}`", "更新地图元素"],
        ["DELETE", "`/map/api/v1/workspaces/{workspace_id}/elements/{element_id}`", "删除地图元素"],
        ["DELETE", "`/map/api/v1/workspaces/{workspace_id}/element-groups/{group_id}/elements`", "按图层删除全部元素"],
    ])
    add_callout(doc, "info", "首屏加载", "终端首次上线后调用「获取图层与元素列表」拉取共享元素；当收到 WebSocket 的图层刷新通知（携带 group_id）时，再次调用该接口刷新对应图层。")
    h2(doc, "6.4 WebSocket 推送")
    add_table(doc, ["方向", "消息", "说明"], [
        ["下行", "图层刷新推送", "某图层多个元素变动时，携带 group_id 通知终端刷新该图层"],
        ["下行", "新增/更新/删除元素推送", "服务端接收到元素变更后，广播给同一 workspace 下各终端"],
    ])

    # ===== 7. 媒体管理 =====
    h1(doc, "7. 媒体管理")
    add_lead(doc, "媒体管理把无人机上的媒体文件（图片/视频）下载到网关本地存储，再上传到对象存储。上传包含自动与手动两种方式，遥控器上云仅支持自动上传。")
    h2(doc, "7.1 上传流程")
    add_bullets(doc, [
        "每次上传前，网关向服务端请求上传临时凭证（storage_config_get）。",
        "网关携带凭证向对象存储执行文件上传，由对象存储校验。",
        "文件传输结束后，网关调用上传结果上报（file_upload_callback）告知服务端结果。",
    ], ordered=True)
    h2(doc, "7.2 媒体 MQTT 方法")
    add_table(doc, ["方法", "Topic 方向", "说明"], [
        ["`storage_config_get`", "requests → requests_reply", "获取上传临时凭证"],
        ["`file_upload_callback`", "events → events_reply", "媒体文件上传结果上报"],
    ])
    add_callout(doc, "info", "说明", "必须是无人机在航点任务时、通过负载控制指令拍摄的媒体文件，才会被媒体管理功能上传归档。")
    h2(doc, "7.3 媒体查询 API")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/media/api/v1/files/{workspace_id}/files`", "媒体文件分页列表"],
        ["GET", "`/media/api/v1/files/{workspace_id}/file/{file_id}/url`", "获取媒体文件短时访问 URL"],
    ])
    add_callout(doc, "warn", "凭据安全", "对象存储访问采用短时授权 URL，不得在日志中记录签名 URL 或长期凭据。")

    # ===== 8. 航线管理 =====
    h1(doc, "8. 航线管理")
    add_lead(doc, "航线管理实现行业作业的批量化、智能化：在云端共享查看、下发执行、取消航线任务并上报进度。航线文件遵循 WPML 规范，一个任务可包含多条航线。")
    h2(doc, "8.1 功能概述")
    add_para(doc, "用户按 WPML 规范编写航线文件（KMZ），上传到平台后创建飞行任务。航线管理引入了「预发布」概念：下发任务给机场与无人机预留准备时间，下发任务后还需调用执行任务。任务进度可实时上报，包含进度与拓展信息。")
    h2(doc, "8.2 任务类型")
    add_table(doc, ["类型", "说明"], [
        ["立即任务", "下发后立即执行；设备端限制 30s 时间误差，收到指令时间与 execute_time 相差超过 30s 将报错"],
        ["定时任务", "指定 execute_time，设备在执行前预留准备时间（如提前 2min 开机搜星、上传航线）"],
        ["条件任务", "需满足 ready_conditions（如电量、存储阈值），全部满足后触发 flighttask_ready 事件"],
    ])
    add_callout(doc, "warn", "注意", "设备正在执行航线任务时，再次收到执行指令将不执行并报错；航线任务指令通常需在无人机关机或停桨时才能执行。")
    h2(doc, "8.3 航线 MQTT 方法")
    add_table(doc, ["方法", "说明"], [
        ["`flighttask_prepare`", "下发（预发布）飞行任务，携带航线与任务参数"],
        ["`flighttask_execute`", "执行已下发的飞行任务"],
        ["`flighttask_progress`", "飞行任务进度与状态上报（事件）"],
        ["`flighttask_ready`", "条件任务就绪事件"],
    ])
    h2(doc, "8.4 航线与任务 API")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/wayline/api/v1/workspaces/{workspace_id}/waylines/file/upload`", "上传 KMZ，字段名 file"],
        ["DELETE", "`/wayline/api/v1/workspaces/{workspace_id}/waylines/{wayline_id}`", "删除航线"],
        ["POST", "`/wayline/api/v1/workspaces/{workspace_id}/flight-tasks`", "创建飞行任务"],
        ["GET", "`/wayline/api/v1/workspaces/{workspace_id}/jobs`", "任务分页"],
        ["PUT", "`/wayline/api/v1/workspaces/{workspace_id}/jobs/{job_id}`", "暂停 / 恢复任务"],
        ["DELETE", "`/wayline/api/v1/workspaces/{workspace_id}/jobs`", "取消任务"],
        ["POST", "`/wayline/api/v1/workspaces/{workspace_id}/jobs/{job_id}/media-highest`", "置顶任务媒体上传优先级"],
    ])
    add_para(doc, "创建任务核心字段：")
    add_code(doc, '{\n  "name": "园区巡检-001",\n  "file_id": "<wayline id>",\n  "dock_sn": "<dock sn>",\n  "task_type": 0,\n  "execute_time": 1785319200000,\n  "rth_altitude": 100,\n  "out_of_control_action": 0,\n  "min_battery_capacity": 60,\n  "min_storage_capacity": 1024\n}')
    add_para(doc, "rth_altitude 返航高度范围 20–500 米，最低电量范围 50–90；定时任务还需 task_days、task_periods。枚举实际值以在线 OpenAPI 与设备能力为准。")

    # ===== 9. HMS 健康管理 =====
    h1(doc, "9. HMS 健康管理")
    add_lead(doc, "HMS（Health Management System）把设备的健康告警显示在指挥中心，提升对风险的感知与处理能力。设备通过事件上报告警，云端拼接文案 Key、查询文案库后渲染展示。")
    h2(doc, "9.1 功能概述")
    add_para(doc, "设备通过健康告警协议（Topic thing/product/{gateway_sn}/events，Method hms）上报机巢与无人机的告警信息到云端。云端拼接文案 Key、查询文案 JSON、渲染文案后在 Web 界面呈现完整告警。")
    add_callout(doc, "info", "全量语义", "HMS 上报的是全量告警信息。若上一次 HMS 上报中的某告警在本次上报中消失，则表示该告警已解除。")
    h2(doc, "9.2 告警文案查询")
    add_table(doc, ["设备类型", "拼接规则", "示例（code=0x1001）"], [
        ["机巢", "`dock_tip_{code}`", "`dock_tip_0x1001`"],
        ["无人机", "`fpv_tip_{code}`", "`fpv_tip_0x0108`"],
        ["遥控器", "`rc_tip_{code}`", "`rc_tip_0x00010001`"],
    ])
    add_para(doc, "其中 {code} 为告警码，通过上报协议字段获得。文案库以部署内置的 hms.json 为准。")
    h2(doc, "9.3 HMS API")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices/hms`", "HMS 告警分页"],
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices/hms/{device_sn}`", "指定设备 HMS 告警"],
        ["PUT", "`/manage/api/v1/devices/{workspace_id}/devices/hms/{device_sn}`", "标记已读 / 更新告警状态"],
    ])
    add_para(doc, "WebSocket 侧以 device_hms 业务码实时推送告警变化。")

    # ===== 10. 指令飞行 =====
    h1(doc, "10. 指令飞行")
    add_lead(doc, "指令飞行解决远程控制中的即时操作限制，让开发者以手动方式继续控制设备或负载。API 分为飞行控制类（DRC）、负载控制类、flyto 指令与一键起飞指令。")
    h2(doc, "10.1 控制权管理")
    add_para(doc, "飞行控制权与负载控制权相互独立，均为排他资源。进入远控前必须先抢占对应控制权，避免多端争抢。")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/control/api/v1/devices/{sn}/authority/flight`", "获取飞行控制权"],
        ["POST", "`/control/api/v1/devices/{sn}/authority/payload`", "获取负载控制权"],
    ])
    h2(doc, "10.2 DRC 指令飞行")
    add_para(doc, "DRC（drone remote control）使用 MQTT 协议并新增上行/下行两个专用 Topic。DRC 指令需先开启指令飞行控制模式（drc_mode_enter）才能使用；其中 drone_control（摇杆控制）必须具备飞行控制权。生命周期：connect → enter → MQTT 心跳/drone_control → exit。")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/control/api/v1/workspaces/{workspace_id}/drc/connect`", "获取 MQTT 短时连接配置"],
        ["POST", "`/control/api/v1/workspaces/{workspace_id}/drc/enter`", "获取控制权并生成精确 Topic ACL"],
        ["POST", "`/control/api/v1/workspaces/{workspace_id}/drc/exit`", "释放控制权并清理 ACL"],
    ])
    add_para(doc, "进入/退出请求：")
    add_code(doc, '{\n  "client_id": "<connect 返回值>",\n  "dock_sn": "<dock sn>",\n  "expire_sec": 3600,\n  "device_info": { "osd_frequency": 10, "hsi_frequency": 1 }\n}')
    add_callout(doc, "warn", "安全联锁", "控制端应每秒发送心跳，并以 10 Hz 发布 drone_control。窗口失焦、按键释放、MQTT 断开或退出时，必须先发送零杆量再释放控制权。急停使用 drone_emergency_stop。")
    add_para(doc, "摇杆控制消息保持设备物模型信封：")
    add_code(doc, '{\n  "tid": "<uuid>", "bid": "<uuid>",\n  "timestamp": 1785319200000,\n  "method": "drone_control",\n  "data": { "seq": 123, "x": 0, "y": 0, "h": 0, "w": 0, "freq": 10, "delay_time": 300 }\n}')
    h2(doc, "10.3 指令控制（flyto / 一键起飞）")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/control/api/v1/devices/{sn}/jobs/{service_identifier}`", "通用服务指令，如返航"],
        ["POST", "`/control/api/v1/devices/{sn}/jobs/fly-to-point`", "飞向目标点（fly_to_point）"],
        ["DELETE", "`/control/api/v1/devices/{sn}/jobs/fly-to-point`", "取消飞向目标点"],
        ["POST", "`/control/api/v1/devices/{sn}/jobs/takeoff-to-point`", "起飞到点（takeoff_to_point）"],
    ])
    h2(doc, "10.4 云台与负载控制")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/control/api/v1/devices/{sn}/payload/commands`", "云台、相机与负载命令"],
    ])
    add_callout(doc, "danger", "能力驱动", "可用的 service_identifier 与负载命令受机型、固件与设备能力限制。前端应按设备能力上报显示，不得把未验证指令作为固定按钮开放。")

    # ===== 11. MQTT Topic 与物模型 =====
    h1(doc, "11. MQTT Topic 与物模型")
    add_lead(doc, "本章汇总设备上云的 MQTT Topic 体系与物模型方法清单，作为各功能的协议底座。")
    h2(doc, "11.1 Topic 定义")
    add_table(doc, ["Topic", "方向", "用途"], [
        ["`sys/product/{gateway_sn}/status`", "设备→云", "上下线状态"],
        ["`thing/product/{gateway_sn}/state`", "设备→云", "属性 / 能力上报"],
        ["`thing/product/{gateway_sn}/events`", "设备→云", "事件上报（HMS、任务进度、上传回调等）"],
        ["`thing/product/{gateway_sn}/events_reply`", "云→设备", "事件应答"],
        ["`thing/product/{gateway_sn}/requests`", "设备→云", "设备发起的请求（如获取凭证）"],
        ["`thing/product/{gateway_sn}/requests_reply`", "云→设备", "请求应答"],
        ["`thing/product/{gateway_sn}/services`", "云→设备", "服务调用下发（如开始直播）"],
        ["`thing/product/{gateway_sn}/services_reply`", "设备→云", "服务调用应答"],
        ["`thing/product/{gateway_sn}/property/set`", "云→设备", "属性设置下发"],
        ["`thing/product/{gateway_sn}/drc/down`", "云→设备", "DRC 下行（drone_control 等）"],
        ["`thing/product/{gateway_sn}/drc/up`", "设备→云", "DRC 上行（响应 / 急停）"],
    ])
    h2(doc, "11.2 物模型方法清单")
    add_table(doc, ["功能", "方法（method）", "类别"], [
        ["设备管理", "update_topo、drone_control_source", "state / 事件"],
        ["直播", "live_capacity、live_start_push、live_stop_push、live_set_quality、live_lens_change", "state / services"],
        ["媒体", "storage_config_get、file_upload_callback", "requests / events"],
        ["航线", "flighttask_prepare、flighttask_execute、flighttask_progress、flighttask_ready", "services / events"],
        ["HMS", "hms", "events"],
        ["指令飞行", "drc_mode_enter、drc_mode_exit、drone_control、drone_emergency_stop、fly_to_point、takeoff_to_point", "services / drc"],
    ])
    add_callout(doc, "info", "ACL", "平台基于 Redis 为每个网关生成精确的 Topic ACL；DRC 进入时按会话生成短时 ACL，退出时清理，避免越权订阅或发布。")

    # ===== 12. HTTPS 接口清单 =====
    h1(doc, "12. HTTPS 接口清单")
    add_lead(doc, "本章汇总平台开放的 HTTP REST 接口，按业务域分组。所有接口默认需携带 x-auth-token。字段级细节以运行实例 OpenAPI（/v3/api-docs、/swagger-ui/index.html）为准。")
    h2(doc, "12.1 鉴权与用户")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/manage/api/v1/login`", "登录（flag=1 Web / flag=2 Pilot）"],
        ["POST", "`/manage/api/v1/token/refresh`", "刷新令牌"],
        ["GET", "`/manage/api/v1/users/current`", "当前用户信息"],
        ["PUT", "`/manage/api/v1/users/current/password`", "修改当前用户密码"],
        ["GET", "`/manage/api/v1/users/{workspace_id}/users`", "工作空间用户列表"],
        ["PUT", "`/manage/api/v1/users/{workspace_id}/users/{user_id}`", "更新用户"],
        ["GET", "`/manage/api/v1/workspaces/current`", "当前工作空间"],
    ])
    h2(doc, "12.2 设备管理")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices`", "设备拓扑列表"],
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices/bound`", "已绑定设备"],
        ["GET", "`/manage/api/v1/devices/{workspace_id}/devices/{device_sn}`", "设备详情"],
        ["POST", "`/manage/api/v1/devices/{device_sn}/binding`", "绑定设备"],
        ["DELETE", "`/manage/api/v1/devices/{device_sn}/unbinding`", "解绑设备"],
        ["PUT", "`/manage/api/v1/devices/{workspace_id}/devices/{device_sn}`", "更新设备信息"],
        ["PUT", "`/manage/api/v1/devices/{workspace_id}/devices/{device_sn}/property`", "设置设备属性"],
        ["POST", "`/manage/api/v1/devices/{workspace_id}/devices/ota`", "下发 OTA 升级"],
    ])
    h2(doc, "12.3 固件管理")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/manage/api/v1/workspaces/firmware-release-notes/latest`", "最新固件发布说明"],
        ["GET", "`/manage/api/v1/workspaces/{workspace_id}/firmwares`", "固件列表"],
        ["POST", "`/manage/api/v1/workspaces/{workspace_id}/firmwares/file/upload`", "上传固件"],
        ["PUT", "`/manage/api/v1/workspaces/{workspace_id}/firmwares/{firmware_id}`", "更新固件信息"],
    ])
    h2(doc, "12.4 直播")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/manage/api/v1/live/capacity`", "直播能力"],
        ["POST", "`/manage/api/v1/live/streams/start`", "开始直播"],
        ["POST", "`/manage/api/v1/live/streams/stop`", "停止直播"],
        ["POST", "`/manage/api/v1/live/streams/update`", "设置清晰度"],
        ["POST", "`/manage/api/v1/live/streams/switch`", "切换镜头"],
    ])
    h2(doc, "12.5 指令飞行与控制权")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/control/api/v1/devices/{sn}/authority/flight`", "获取飞行控制权"],
        ["POST", "`/control/api/v1/devices/{sn}/authority/payload`", "获取负载控制权"],
        ["POST", "`/control/api/v1/devices/{sn}/jobs/{service_identifier}`", "通用服务指令"],
        ["POST", "`/control/api/v1/devices/{sn}/jobs/fly-to-point`", "飞向目标点"],
        ["DELETE", "`/control/api/v1/devices/{sn}/jobs/fly-to-point`", "取消飞向目标点"],
        ["POST", "`/control/api/v1/devices/{sn}/jobs/takeoff-to-point`", "起飞到点"],
        ["POST", "`/control/api/v1/devices/{sn}/payload/commands`", "负载 / 云台命令"],
        ["POST", "`/control/api/v1/workspaces/{workspace_id}/drc/connect`", "DRC 连接配置"],
        ["POST", "`/control/api/v1/workspaces/{workspace_id}/drc/enter`", "DRC 进入"],
        ["POST", "`/control/api/v1/workspaces/{workspace_id}/drc/exit`", "DRC 退出"],
    ])
    h2(doc, "12.6 航线与任务")
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "`/wayline/api/v1/workspaces/{workspace_id}/waylines/file/upload`", "上传 KMZ"],
        ["DELETE", "`/wayline/api/v1/workspaces/{workspace_id}/waylines/{wayline_id}`", "删除航线"],
        ["POST", "`/wayline/api/v1/workspaces/{workspace_id}/flight-tasks`", "创建飞行任务"],
        ["GET", "`/wayline/api/v1/workspaces/{workspace_id}/jobs`", "任务分页"],
        ["PUT", "`/wayline/api/v1/workspaces/{workspace_id}/jobs/{job_id}`", "暂停 / 恢复"],
        ["DELETE", "`/wayline/api/v1/workspaces/{workspace_id}/jobs`", "取消任务"],
        ["POST", "`/wayline/api/v1/workspaces/{workspace_id}/jobs/{job_id}/media-highest`", "媒体优先上传"],
    ])
    h2(doc, "12.7 地图元素与飞行区")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/map/api/v1/workspaces/{workspace_id}/element-groups`", "图层与元素列表"],
        ["POST", "`/map/api/v1/workspaces/{workspace_id}/element-groups/{group_id}/elements`", "创建元素"],
        ["PUT", "`/map/api/v1/workspaces/{workspace_id}/elements/{element_id}`", "更新元素"],
        ["DELETE", "`/map/api/v1/workspaces/{workspace_id}/elements/{element_id}`", "删除元素"],
        ["GET", "`/map/api/v1/workspaces/{workspace_id}/device-status`", "设备态势状态"],
        ["GET", "`/map/api/v1/workspaces/{workspace_id}/flight-areas`", "飞行区列表"],
        ["POST", "`/map/api/v1/workspaces/{workspace_id}/flight-area`", "创建飞行区"],
        ["PUT", "`/map/api/v1/workspaces/{workspace_id}/flight-area/{area_id}`", "更新飞行区"],
        ["DELETE", "`/map/api/v1/workspaces/{workspace_id}/flight-area/{area_id}`", "删除飞行区"],
        ["POST", "`/map/api/v1/workspaces/{workspace_id}/flight-area/sync`", "同步飞行区至设备"],
    ])
    h2(doc, "12.8 媒体、日志与运维")
    add_table(doc, ["方法", "路径", "说明"], [
        ["GET", "`/media/api/v1/files/{workspace_id}/files`", "媒体文件列表"],
        ["GET", "`/media/api/v1/files/{workspace_id}/file/{file_id}/url`", "媒体短时 URL"],
        ["GET", "`/manage/api/v1/workspaces/{workspace_id}/devices/{device_sn}/logs-uploaded`", "已上传日志"],
        ["POST", "`/manage/api/v1/workspaces/{workspace_id}/devices/{device_sn}/logs`", "发起日志上传"],
        ["GET", "`/manage/api/v1/workspaces/{workspace_id}/logs/{logs_id}/url/{file_id}`", "日志文件下载 URL"],
        ["GET", "`/manage/api/v1/ops/logs`", "操作 / 交互日志"],
    ])

    # ===== 13. WebSocket 接口 =====
    h1(doc, "13. WebSocket 接口")
    add_lead(doc, "平台通过单一 WebSocket 通道向前端推送设备状态、遥测、任务进度、告警与交互日志。前端按 biz_code 分发处理。")
    h2(doc, "13.1 连接与认证")
    add_code(doc, "ws(s)://<host>:<pilot-port>/api/v1/ws?x-auth-token=<access_token>")
    add_para(doc, "消息统一结构：")
    add_code(doc, '{\n  "biz_code": "device_osd",\n  "version": "1.0",\n  "timestamp": 1785319200000,\n  "data": { }\n}')
    h2(doc, "13.2 业务码（biz_code）")
    add_table(doc, ["biz_code", "说明"], [
        ["`device_update_topo`", "设备上线/下线/拓扑更新"],
        ["`device_osd`", "飞行器遥测 OSD"],
        ["`dock_osd`", "机巢遥测 OSD"],
        ["`gateway_osd`", "网关遥测 OSD"],
        ["`live_capacity`", "直播能力更新"],
        ["`live_status`", "直播状态变化"],
        ["`control_source_change`", "控制权来源变化"],
        ["`payload_authority_grab`", "负载控制权抢占"],
        ["`device_hms`", "HMS 健康告警"],
        ["`*_progress`", "任务进度类（如 flighttask 进度）"],
        ["`interaction_log`", "交互 / 操作日志"],
    ])

    # ===== 14. 航线文件格式标准 WPML =====
    h1(doc, "14. 航线文件格式标准（WPML）")
    add_lead(doc, "WPML（WayPoint Markup Language）是基于 KML 扩展的航线文件格式标准，以 KMZ 压缩包承载，用于描述航线模板与可执行航线。")
    h2(doc, "14.1 什么是 WPML")
    add_bullets(doc, [
        "开放：定义公开共享，采用文件包归档、标记语言描述路径，可扩展性强。",
        "易读：使用基础 XML 与媒体文件格式，可不依赖解析软件即可读写。",
        "继承：作为航线数字资产载体持续迭代并保持代际兼容，可在不同设备与机型间应用。",
    ])
    h2(doc, "14.2 KMZ 文件结构")
    add_table(doc, ["文件", "作用"], [
        ["`wpmz/template.kml`", "模板文件：定义业务属性（作业类型、坐标/高度参考、全局速度、航点与动作模板）"],
        ["`wpmz/waylines.wpml`", "执行文件：由模板解算出的、设备可直接执行的具体航点与动作序列"],
        ["`wpmz/res/`", "资源目录：航线所需辅助资源（如精准复拍参考目标物照片）"],
    ])
    add_code(doc, "mission.kmz\n└── wpmz/\n    ├── template.kml\n    ├── waylines.wpml\n    └── res/")
    add_callout(doc, "warn", "命名规范", "航线文件内部各文件/文件夹命名须遵循此规范，否则可能读取失败。航线名称即 KMZ 文件名。")
    h2(doc, "14.3 WPML 核心元素")
    add_para(doc, "WPML 使用 <wpml:*> 命名空间扩展标准 KML，元素分为航线级、航点级、动作级三层。")
    add_bullets(doc, [
        "航线级：templateType 作业类型、waylineCoordinateSysParam 坐标/高度、autoFlightSpeed 速度、globalRTHHeight 返航高度、missionConfig 任务全局配置。",
        "航点级：Placemark+Point/coordinates 坐标、index 序号、height/ellipsoidHeight/executeHeight 高度、waypointHeadingParam 航向、waypointTurnParam 转弯。",
        "动作级：actionGroup 动作组、actionTrigger 触发器、action+actionActuatorFunc 具体动作。",
    ])
    add_callout(doc, "info", "模板与执行分离", "template.kml 描述「意图」，waylines.wpml 描述「执行」。航线规划软件读取模板参数、结合路径生成算法解算出执行航线。倾斜摄影模板会生成 5 条可执行航线（1 正射 + 4 倾斜）。")

    # ===== 15. WPML 元素字段参考 =====
    h1(doc, "15. WPML 元素字段参考")
    add_lead(doc, "本章列出 WPML 常用字段，按任务配置、坐标与高度、航点、动作触发与动作类型分组。字段的机型适配与取值范围以目标设备能力为准。")
    h2(doc, "15.1 任务配置（wpml:missionConfig）")
    add_table(doc, ["元素", "说明", "取值 / 枚举"], [
        ["`wpml:flyToWaylineMode`", "飞向首航点模式", "safely / pointToPoint"],
        ["`wpml:finishAction`", "航线结束动作", "goHome / noAction / autoLand / gotoFirstWaypoint"],
        ["`wpml:exitOnRCLost`", "失控是否继续执行航线", "goContinue / executeLostAction"],
        ["`wpml:executeRCLostAction`", "失控动作类型", "goBack / landing / hover"],
        ["`wpml:takeOffSecurityHeight`", "安全起飞高度", "浮点，米（相对起飞点）"],
        ["`wpml:globalTransitionalSpeed`", "全局航线过渡速度", "[1,15] m/s"],
        ["`wpml:globalRTHHeight`", "全局返航高度", "浮点，米"],
        ["`wpml:droneInfo / payloadInfo`", "无人机 / 负载机型信息", "主类型+子类型+挂载位置"],
    ])
    h2(doc, "15.2 坐标系与高度（wpml:waylineCoordinateSysParam）")
    add_table(doc, ["元素", "说明", "取值 / 枚举"], [
        ["`wpml:coordinateMode`", "经纬度坐标系", "WGS84"],
        ["`wpml:heightMode`", "高程参考平面", "EGM96 / relativeToStartPoint / aboveGroundLevel / realTimeFollowSurface"],
        ["`wpml:positioningType`", "定位数据源", "GPS / RTKBaseStation / QianXun / Custom"],
        ["`wpml:globalShootHeight`", "离被摄面高度（测绘）", "浮点，米"],
        ["`wpml:surfaceFollowModeEnable`", "是否开启仿地飞行", "0 / 1"],
        ["`wpml:executeHeightMode`", "执行高度模式（仅 waylines.wpml）", "WGS84 / relativeToStartPoint"],
    ])
    h2(doc, "15.3 航点（Placemark）")
    add_table(doc, ["元素", "说明", "取值 / 枚举"], [
        ["`Point/coordinates`", "航点经纬度", "经度,纬度[,高度]"],
        ["`wpml:index`", "航点序号", "[0,65535]，从 0 递增"],
        ["`wpml:executeHeight`", "航点执行高度", "浮点，米"],
        ["`wpml:waypointSpeed`", "航点飞行速度", "[1,15] m/s"],
        ["`wpml:waypointHeadingMode`", "偏航角模式", "followWayline / manually / fixed / smoothTransition / towardPOI"],
        ["`wpml:waypointTurnMode`", "转弯模式", "coordinateTurn / toPointAndStop... / toPointAndPass..."],
        ["`wpml:gimbalPitchAngle`", "航点云台俯仰角", "浮点，°（机型范围）"],
    ])
    h2(doc, "15.4 动作触发（wpml:actionTrigger）")
    add_table(doc, ["触发类型", "说明"], [
        ["`reachPoint`", "到达航点时执行"],
        ["`betweenAdjacentPoints`", "航段触发（配合 gimbalEvenlyRotate 均匀转云台）"],
        ["`multipleTiming`", "等时触发（参数为间隔秒）"],
        ["`multipleDistance`", "等距触发（参数为间隔米）"],
    ])
    h2(doc, "15.5 动作类型（wpml:actionActuatorFunc）")
    add_table(doc, ["动作", "说明", "关键参数"], [
        ["`takePhoto`", "单拍", "payloadPositionIndex、payloadLensIndex、fileSuffix"],
        ["`startRecord / stopRecord`", "开始 / 结束录像", "payloadPositionIndex、payloadLensIndex"],
        ["`focus`", "对焦", "isPointFocus、focusX/Y、focusRegionWidth/Height"],
        ["`zoom`", "变焦", "focalLength（mm）"],
        ["`gimbalRotate`", "旋转云台", "gimbalPitch/Roll/YawRotateAngle、gimbalRotateMode"],
        ["`gimbalEvenlyRotate`", "航段间均匀转动云台 pitch", "gimbalPitchRotateAngle"],
        ["`rotateYaw`", "无人机偏航", "aircraftHeading、aircraftPathMode"],
        ["`hover`", "悬停等待", "hoverTime（s）"],
        ["`customDirName`", "创建新文件夹", "directoryName"],
        ["`orientedShoot / accurateShoot`", "定向 / 精准复拍", "相机姿态、目标框、文件信息等"],
        ["`openAIDetect`", "AI 目标检测", "aiDetectObjectType、aiDetectObjectAction"],
    ])
    add_callout(doc, "info", "存储类型", "拍照/录像的镜头存储类型（payloadLensIndex）常见枚举：wide 广角、zoom 变焦、ir 红外、narrow_band 窄带、visable 可见光；可组合，如 wide,ir。具体支持项以负载能力为准。")

    # ===== 16. 错误码与安全基线 =====
    h1(doc, "16. 错误码与安全基线")
    add_lead(doc, "本章说明统一错误处理约定与部署安全基线。字段级错误码以运行实例 OpenAPI 与设备返回为准。")
    h2(doc, "16.1 错误处理约定")
    add_table(doc, ["层级", "取值", "含义"], [
        ["业务 code", "`0`", "成功"],
        ["业务 code / HTTP", "`401`", "令牌无效或过期，需重新登录 / 刷新"],
        ["HTTP", "`400`", "请求参数错误"],
        ["HTTP", "`403`", "无权限（工作空间 / 控制权 / ACL）"],
        ["HTTP", "`404`", "资源不存在"],
        ["HTTP", "`5xx`", "服务端异常"],
        ["业务 code", "非 0", "业务失败，message 描述原因"],
    ])
    add_callout(doc, "info", "设备侧错误", "MQTT 服务调用与 DRC 指令的错误通过 *_reply 应答中的结果码返回；HMS 告警码通过文案 Key 映射为可读文案。")
    h2(doc, "16.2 安全基线")
    add_callout(doc, "danger", "上线前必做", "修改初始化密码与所有示例密钥；启用 TLS/WSS；限制 MQTT/RTSP 源 IP；管理端口不出公网；不得在日志中记录 JWT、MQTT 密码或对象存储签名 URL。")
    add_bullets(doc, [
        "令牌：JWT 短时有效，DRC 连接密码为一次性短时凭据，客户端不得持久化。",
        "ACL：按网关与会话生成精确 Topic ACL，DRC 退出时清理。",
        "能力驱动：所有指令按设备能力上报开放，不下发未验证指令。",
        "控制权联锁：释放控制权前先发送零杆量，断连即安全停止。",
    ])


if __name__ == "__main__":
    build()
